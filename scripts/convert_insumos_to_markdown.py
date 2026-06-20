#!/usr/bin/env python3
"""Convert source documents from 01-insumos_originales to Markdown."""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import sys
import textwrap
from dataclasses import dataclass
from datetime import date
from html.parser import HTMLParser
from pathlib import Path
from typing import Callable


ROOT = Path(__file__).resolve().parents[1]
INPUT_DIR = ROOT / "01-insumos_originales"
OUTPUT_DIR = ROOT / "02-material_de_referencia"
INDEX_PATH = OUTPUT_DIR / "README.md"
MANIFEST_PATH = OUTPUT_DIR / "conversion_manifest.json"

SUPPORTED_FORMATS = {".txt", ".md", ".html", ".htm", ".docx", ".pdf"}
PDF_MIN_TEXT_CHARS = 500
PDF_MIN_AVG_CHARS_PER_PAGE = 80

AUTO_NOTE = (
    "> Documento convertido automáticamente desde el insumo original. Revisar "
    "estructura, citas y posibles errores de extracción antes de usarlo como "
    "fuente docente."
)


@dataclass
class ConversionResult:
    source: Path
    output: Path | None
    formato: str
    estado: str
    hash_origen: str = ""
    observaciones: str = ""
    error: str = ""


class SimpleMarkdownHTMLParser(HTMLParser):
    """Small HTML to Markdown-ish text parser used when BeautifulSoup is absent."""

    BLOCK_TAGS = {
        "address",
        "article",
        "aside",
        "blockquote",
        "br",
        "div",
        "footer",
        "header",
        "li",
        "main",
        "nav",
        "ol",
        "p",
        "pre",
        "section",
        "table",
        "tbody",
        "td",
        "tfoot",
        "th",
        "thead",
        "tr",
        "ul",
    }

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.skip_depth = 0
        self.heading_level: int | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript"}:
            self.skip_depth += 1
            return
        if self.skip_depth:
            return
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self._newline(2)
            self.heading_level = int(tag[1])
            self.parts.append("#" * self.heading_level + " ")
        elif tag == "li":
            self._newline(1)
            self.parts.append("- ")
        elif tag in self.BLOCK_TAGS:
            self._newline(1 if tag == "br" else 2)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript"} and self.skip_depth:
            self.skip_depth -= 1
            return
        if self.skip_depth:
            return
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self.heading_level = None
            self._newline(2)
        elif tag in self.BLOCK_TAGS:
            self._newline(2)

    def handle_data(self, data: str) -> None:
        if self.skip_depth:
            return
        text = " ".join(html.unescape(data).split())
        if not text:
            return
        if self.parts and not self.parts[-1].endswith(("\n", " ", "# ", "- ")):
            self.parts.append(" ")
        self.parts.append(text)

    def get_markdown(self) -> str:
        return normalize_markdown("".join(self.parts))

    def _newline(self, count: int) -> None:
        current = "".join(self.parts[-3:])
        existing = len(current) - len(current.rstrip("\n"))
        if existing < count:
            self.parts.append("\n" * (count - existing))


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Convierte documentos de 01-insumos_originales a Markdown."
    )
    parser.add_argument("--input-dir", type=Path, default=INPUT_DIR)
    parser.add_argument("--output-dir", type=Path, default=OUTPUT_DIR)
    args = parser.parse_args()

    input_dir = args.input_dir.resolve()
    output_dir = args.output_dir.resolve()

    if not input_dir.exists():
        print(f"No existe el directorio de insumos: {input_dir}", file=sys.stderr)
        return 1

    output_dir.mkdir(parents=True, exist_ok=True)

    sources = sorted(path for path in input_dir.iterdir() if path.is_file())
    index_path = output_dir / "README.md"
    manifest_path = output_dir / "conversion_manifest.json"
    manifest_entries = load_manifest(manifest_path)
    known_conversions = collect_known_conversions(output_dir, index_path, manifest_entries)
    results: list[ConversionResult] = []

    for source in sources:
        formato = source.suffix.lower().lstrip(".") or "sin_extension"
        source_rel = relative_or_absolute(source)
        source_hash = file_sha256(source)

        if source.suffix.lower() not in SUPPORTED_FORMATS:
            result = ConversionResult(
                source=source,
                output=None,
                formato=formato,
                estado="omitido_formato_no_soportado",
                hash_origen=source_hash,
                observaciones="Formato no soportado todavía.",
            )
            results.append(result)
            upsert_manifest_entry(manifest_entries, result)
            continue

        known = known_conversions.get(source_rel)
        if (
            known
            and known.estado in {"convertido", "omitido_ya_convertido"}
            and known.hash_origen in {"", source_hash}
            and (known.output is None or known.output.exists())
        ):
            result = ConversionResult(
                source=source,
                output=known.output,
                formato=formato,
                estado="omitido_ya_convertido",
                hash_origen=source_hash,
                observaciones="Ya existe conversión previa.",
            )
            results.append(result)
            upsert_manifest_entry(manifest_entries, result)
            continue

        observation_prefix = ""
        versioned = False
        if (
            known
            and known.estado in {"convertido", "omitido_ya_convertido"}
            and known.hash_origen
            and known.hash_origen != source_hash
        ):
            observation_prefix = (
                "Archivo original con nombre previamente convertido, pero contenido distinto."
            )
            versioned = True

        output_path = allocate_new_output_path(source, output_dir, versioned=versioned)

        try:
            body, observation = extract_markdown(source)
            observation = join_observations(observation_prefix, observation)
            markdown = build_markdown(source, formato, source_hash, body, observation)
            output_path.write_text(markdown, encoding="utf-8")
            result = ConversionResult(
                source=source,
                output=output_path,
                formato=formato,
                estado="convertido",
                hash_origen=source_hash,
                observaciones=observation,
            )
            results.append(result)
            known_conversions[source_rel] = KnownConversion(
                output=output_path,
                hash_origen=source_hash,
                source="current_run",
            )
            upsert_manifest_entry(manifest_entries, result)
        except Exception as exc:  # Keep processing the rest of the batch.
            result = ConversionResult(
                source=source,
                output=output_path,
                formato=formato,
                estado="error",
                hash_origen=source_hash,
                observaciones=observation_prefix,
                error=str(exc),
            )
            results.append(result)
            upsert_manifest_entry(manifest_entries, result)

    write_index(index_path, results, input_dir, output_dir)
    write_manifest(manifest_path, manifest_entries)
    print_summary(results, index_path, manifest_path)
    return 0


def sanitize_stem(stem: str) -> str:
    cleaned = re.sub(r"\s+", " ", stem.strip())
    return cleaned or "documento"


@dataclass
class KnownConversion:
    output: Path | None
    hash_origen: str = ""
    source: str = ""
    estado: str = ""


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_manifest(manifest_path: Path) -> list[dict[str, str]]:
    if not manifest_path.exists():
        return []
    try:
        raw = json.loads(manifest_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return []
    if isinstance(raw, list):
        return [entry for entry in raw if isinstance(entry, dict)]
    if isinstance(raw, dict) and isinstance(raw.get("entries"), list):
        return [entry for entry in raw["entries"] if isinstance(entry, dict)]
    return []


def write_manifest(manifest_path: Path, entries: list[dict[str, str]]) -> None:
    ordered = sorted(entries, key=lambda item: item.get("archivo_origen", ""))
    payload = {
        "version": 1,
        "fecha_actualizacion": date.today().isoformat(),
        "entries": ordered,
    }
    manifest_path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def upsert_manifest_entry(entries: list[dict[str, str]], result: ConversionResult) -> None:
    source_rel = relative_or_absolute(result.source)
    output_rel = relative_or_absolute(result.output) if result.output else ""
    payload = {
        "archivo_origen": source_rel,
        "markdown_generado": output_rel,
        "formato_origen": result.formato,
        "hash_origen": result.hash_origen,
        "fecha_conversion": date.today().isoformat(),
        "estado": result.estado,
        "observaciones": result.error or result.observaciones,
    }
    for index, entry in enumerate(entries):
        if entry.get("archivo_origen") == source_rel:
            entries[index] = {**entry, **payload}
            return
    entries.append(payload)


def collect_known_conversions(
    output_dir: Path, index_path: Path, manifest_entries: list[dict[str, str]]
) -> dict[str, KnownConversion]:
    known: dict[str, KnownConversion] = {}
    for entry in manifest_entries:
        source = entry.get("archivo_origen", "")
        if not source:
            continue
        output = path_from_repo(entry.get("markdown_generado", ""))
        known[source] = KnownConversion(
            output=output,
            hash_origen=entry.get("hash_origen", ""),
            source="manifest",
            estado=entry.get("estado", ""),
        )

    for md_path in sorted(output_dir.glob("*.md")):
        if md_path.name == "README.md":
            continue
        frontmatter = read_frontmatter(md_path)
        source = frontmatter.get("archivo_origen", "")
        if source and source not in known:
            known[source] = KnownConversion(
                output=md_path,
                hash_origen=frontmatter.get("hash_origen", ""),
                source="frontmatter",
                estado="convertido",
            )

    for source, output in read_index_conversions(index_path).items():
        if source not in known:
            known[source] = KnownConversion(output=output, source="index", estado="convertido")

    input_dir = ROOT / "01-insumos_originales"
    for source in input_dir.iterdir() if input_dir.exists() else []:
        if not source.is_file():
            continue
        source_rel = relative_or_absolute(source)
        same_base = output_dir / f"{sanitize_stem(source.stem)}.md"
        if source_rel not in known and same_base.exists():
            known[source_rel] = KnownConversion(
                output=same_base, source="same_base", estado="convertido"
            )
    return known


def read_frontmatter(path: Path) -> dict[str, str]:
    text = path.read_text(encoding="utf-8", errors="replace")
    match = re.match(r"---\n(.*?)\n---", text, flags=re.DOTALL)
    if not match:
        return {}
    data: dict[str, str] = {}
    for line in match.group(1).splitlines():
        key, sep, value = line.partition(":")
        if not sep:
            continue
        value = value.strip()
        if len(value) >= 2 and value[0] == value[-1] == '"':
            value = value[1:-1].replace('\\"', '"').replace("\\\\", "\\")
        data[key.strip()] = value
    return data


def read_index_conversions(index_path: Path) -> dict[str, Path | None]:
    if not index_path.exists():
        return {}
    conversions: dict[str, Path | None] = {}
    row_pattern = re.compile(r"^\|\s*`([^`]+)`\s*\|\s*(?:`([^`]+)`)?")
    for line in index_path.read_text(encoding="utf-8", errors="replace").splitlines():
        match = row_pattern.match(line)
        if not match:
            continue
        source = match.group(1)
        output = path_from_repo(match.group(2) or "")
        conversions[source] = output
    return conversions


def path_from_repo(value: str) -> Path | None:
    if not value:
        return None
    path = Path(value)
    return path if path.is_absolute() else ROOT / path


def allocate_new_output_path(source: Path, output_dir: Path, versioned: bool = False) -> Path:
    stem = sanitize_stem(source.stem)
    if versioned:
        for index in range(2, 1000):
            candidate = output_dir / f"{stem}_v{index}.md"
            if not candidate.exists():
                return candidate
    candidate = output_dir / f"{stem}.md"
    if not candidate.exists():
        return candidate
    for index in range(2, 1000):
        candidate = output_dir / f"{stem}_{index:02d}.md"
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"No se pudo asignar nombre de salida para {source.name}")


def join_observations(*parts: str) -> str:
    return " ".join(part.strip() for part in parts if part and part.strip())


def extract_markdown(source: Path) -> tuple[str, str]:
    suffix = source.suffix.lower()
    extractors: dict[str, Callable[[Path], tuple[str, str]]] = {
        ".txt": extract_plain_text,
        ".md": extract_plain_text,
        ".html": extract_html,
        ".htm": extract_html,
        ".docx": extract_docx,
        ".pdf": extract_pdf,
    }
    return extractors[suffix](source)


def extract_plain_text(source: Path) -> tuple[str, str]:
    text = source.read_text(encoding="utf-8", errors="replace")
    return normalize_markdown(text), ""


def extract_html(source: Path) -> tuple[str, str]:
    raw = source.read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        lines: list[str] = []
        for element in soup.find_all(
            ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "blockquote"]
        ):
            text = " ".join(element.get_text(" ", strip=True).split())
            if not text:
                continue
            tag_name = element.name.lower()
            if tag_name.startswith("h") and tag_name[1:].isdigit():
                level = min(int(tag_name[1:]), 6)
                lines.append(f"{'#' * level} {text}")
            elif tag_name == "li":
                lines.append(f"- {text}")
            elif tag_name == "blockquote":
                lines.append(f"> {text}")
            else:
                lines.append(text)

        if lines:
            return normalize_markdown("\n\n".join(lines)), ""
    except ImportError:
        pass

    fallback = SimpleMarkdownHTMLParser()
    fallback.feed(raw)
    return fallback.get_markdown(), "HTML convertido con parser basico de Python."


def extract_docx(source: Path) -> tuple[str, str]:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "Falta la dependencia python-docx. Ejecuta: pip install -r requirements.txt"
        ) from exc

    document = docx.Document(source)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            level_match = re.search(r"(\d+)", style_name)
            level = min(int(level_match.group(1)), 6) if level_match else 2
            blocks.append(f"{'#' * level} {text}")
        else:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    return normalize_markdown("\n\n".join(blocks)), ""


def extract_pdf(source: Path) -> tuple[str, str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "Falta la dependencia pypdf. Ejecuta: pip install -r requirements.txt"
        ) from exc

    reader = PdfReader(str(source))
    pages_text: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = normalize_pdf_text(text)
        if text:
            pages_text.append(text)
        else:
            pages_text.append(f"<!-- Pagina {index}: sin texto extraido -->")

    body = normalize_markdown("\n\n".join(pages_text))
    extracted_chars = len(re.sub(r"\s+", "", body))
    page_count = max(len(reader.pages), 1)
    observation = ""
    if (
        extracted_chars < PDF_MIN_TEXT_CHARS
        or extracted_chars / page_count < PDF_MIN_AVG_CHARS_PER_PAGE
    ):
        observation = (
            "Posible PDF escaneado o extracción insuficiente; requiere OCR o "
            "revisión manual."
        )
    return body, observation


def normalize_pdf_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
    paragraphs: list[str] = []
    current: list[str] = []

    for line in lines:
        if not line:
            if current:
                paragraphs.append(" ".join(current))
                current = []
            continue
        current.append(line)

    if current:
        paragraphs.append(" ".join(current))

    return "\n\n".join(paragraphs)


def normalize_markdown(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [line.rstrip() for line in text.split("\n")]
    return "\n".join(lines).strip()


def build_markdown(
    source: Path, formato: str, source_hash: str, body: str, observation: str
) -> str:
    title = sanitize_title(source.stem)
    frontmatter = {
        "titulo": title,
        "archivo_origen": source.relative_to(ROOT).as_posix(),
        "formato_origen": formato,
        "hash_origen": source_hash,
        "fecha_conversion": date.today().isoformat(),
        "estado_conversion": "borrador",
        "observaciones": observation,
    }
    yaml = "\n".join(
        [
            "---",
            f'titulo: "{yaml_escape(frontmatter["titulo"])}"',
            f'archivo_origen: "{yaml_escape(frontmatter["archivo_origen"])}"',
            f'formato_origen: "{yaml_escape(frontmatter["formato_origen"])}"',
            f'hash_origen: "{frontmatter["hash_origen"]}"',
            f'fecha_conversion: "{frontmatter["fecha_conversion"]}"',
            f'estado_conversion: "{frontmatter["estado_conversion"]}"',
            f'observaciones: "{yaml_escape(frontmatter["observaciones"])}"',
            "---",
        ]
    )
    body = body or "_No se extrajo texto del documento original._"
    return f"{yaml}\n\n{AUTO_NOTE}\n\n{body}\n"


def sanitize_title(stem: str) -> str:
    return re.sub(r"\s+", " ", stem.strip()) or "Documento sin titulo"


def yaml_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace('"', '\\"')


def write_index(
    index_path: Path, results: list[ConversionResult], input_dir: Path, output_dir: Path
) -> None:
    converted = sum(1 for result in results if result.estado == "convertido")
    already = sum(1 for result in results if result.estado == "omitido_ya_convertido")
    unsupported = sum(
        1 for result in results if result.estado == "omitido_formato_no_soportado"
    )
    errors = sum(1 for result in results if result.estado == "error")

    lines = [
        "# Índice de conversión",
        "",
        f"Fecha de conversión: {date.today().isoformat()}",
        "",
        f"Directorio de origen: `{relative_or_absolute(input_dir)}`",
        f"Directorio de salida: `{relative_or_absolute(output_dir)}`",
        "",
        "## Resumen",
        "",
        f"- Archivos detectados: {len(results)}",
        f"- Ya convertidos: {already}",
        f"- Nuevos convertidos: {converted}",
        f"- Omitidos por formato no soportado: {unsupported}",
        f"- Archivos con error: {errors}",
        "",
        "## Detalle",
        "",
        "| Archivo original | Markdown generado | Formato | Hash corto | Estado | Observaciones / errores |",
        "| --- | --- | --- | --- | --- | --- |",
    ]

    for result in results:
        source = relative_or_absolute(result.source)
        output = relative_or_absolute(result.output) if result.output else ""
        note = result.error or result.observaciones
        short_hash = result.hash_origen[:12] if result.hash_origen else ""
        lines.append(
            "| "
            + " | ".join(
                [
                    table_escape(f"`{source}`"),
                    table_escape(f"`{output}`" if output else ""),
                    table_escape(result.formato),
                    table_escape(short_hash),
                    table_escape(result.estado),
                    table_escape(note),
                ]
            )
            + " |"
        )

    index_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def relative_or_absolute(path: Path | None) -> str:
    if path is None:
        return ""
    try:
        return path.resolve().relative_to(ROOT).as_posix()
    except ValueError:
        return path.as_posix()


def table_escape(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", " ")


def print_summary(
    results: list[ConversionResult], index_path: Path, manifest_path: Path
) -> None:
    converted = sum(1 for result in results if result.estado == "convertido")
    already = sum(1 for result in results if result.estado == "omitido_ya_convertido")
    unsupported = sum(
        1 for result in results if result.estado == "omitido_formato_no_soportado"
    )
    errors = sum(1 for result in results if result.estado == "error")
    message = f"""
    Archivos detectados: {len(results)}
    Ya convertidos: {already}
    Nuevos convertidos: {converted}
    Omitidos por formato no soportado: {unsupported}
    Errores: {errors}
    Índice actualizado: {relative_or_absolute(index_path)}
    Manifiesto actualizado: {relative_or_absolute(manifest_path)}
    """
    print(textwrap.dedent(message).strip())


if __name__ == "__main__":
    raise SystemExit(main())
